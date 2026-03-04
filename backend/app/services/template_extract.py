"""Extract text and section headings from uploaded template documents."""
import re
from pathlib import Path


def extract_text_from_file(file_path: Path, content_type: str | None = None) -> str:
    """Extract plain text from .docx or .txt. Returns empty string on failure."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception:
            return ""
    return ""


def _is_heading_style(style_name: str | None) -> bool:
    if not style_name:
        return False
    n = (style_name or "").lower()
    return (
        n.startswith("heading")
        or n == "title"
        or n == "subtitle"
    )


def extract_headings_from_file(file_path: Path) -> list[str]:
    """Extract section titles/headings from .docx (by style) or .txt (heuristic). Returns list of strings."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            titles = []
            for p in doc.paragraphs:
                style_name = getattr(p.style, "name", None) if p.style else None
                if _is_heading_style(style_name) and (p.text or "").strip():
                    titles.append((p.text or "").strip())
            return titles
        except Exception:
            return []
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="replace")
        return _extract_headings_from_plain_text(text)
    return []


def extract_headings_from_text(text: str) -> list[str]:
    """Extract section titles from plain text (e.g. after pulling from Word). Use as fallback when style-based extraction returns nothing."""
    if not (text or "").strip():
        return []
    return _extract_headings_from_plain_text(text.strip())


def _extract_headings_from_plain_text(text: str) -> list[str]:
    """Heuristic: lines that look like section headers (numbered, ALL CAPS, or short title-like)."""
    lines = [ln.strip() for ln in text.splitlines()]
    seen = set()
    titles = []
    for i, line in enumerate(lines):
        if not line or len(line) > 70:
            continue
        cand = None
        # Numbered section: "1. Overview", "1) Attendees"
        if re.match(r"^\d+[.)]\s+.+", line):
            cand = re.sub(r"^\d+[.)]\s*", "", line).strip()
        # ALL CAPS line (common for section headers)
        elif line.isupper() and 2 < len(line) <= 60:
            cand = line
        # Short line, no period — section break (next blank) or likely heading (short line followed by content)
        elif 3 <= len(line) <= 45 and not line.endswith(".") and "," not in line:
            next_blank = i + 1 >= len(lines) or not lines[i + 1]
            next_is_content = (
                i + 1 < len(lines)
                and lines[i + 1]
                and len(lines[i + 1]) > len(line)
                and len(line) <= 35
            )
            if next_blank or next_is_content:
                cand = line
        if cand and cand.lower() not in seen:
            seen.add(cand.lower())
            titles.append(cand)
    return titles[:25]
